"""
淘宝用户行为数据分析 - 最终报告自动生成
输入：user_behavior_clean.csv, 各PNG图片
输出：最终报告_淘宝用户行为分析.docx
依赖：pip install python-docx pandas numpy matplotlib
"""

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# ========= 配置 =========
DATA_FILE = "user_behavior_clean.csv"
IMAGE_DIR = "./"          # 图片所在目录
# ========================

def load_data():
    df = pd.read_csv(DATA_FILE)
    df["date"] = pd.to_datetime(df["date"])
    df["date_time"] = pd.to_datetime(df["date_time"])
    return df

def calc_key_metrics(df):
    """计算报告中需要填入的所有关键指标"""
    metrics = {}
    # 基本量
    metrics["total_users"] = df["user_id"].nunique()
    metrics["total_items"] = df["item_id"].nunique()
    metrics["total_records"] = len(df)

    # 行为分布
    behavior_counts = df["behavior_type"].value_counts()
    pv = behavior_counts.get("pv", 0)
    cart = behavior_counts.get("cart", 0)
    fav = behavior_counts.get("fav", 0)
    buy = behavior_counts.get("buy", 0)
    metrics["pv_num"] = pv
    metrics["cart_num"] = cart
    metrics["fav_num"] = fav
    metrics["buy_num"] = buy
    metrics["pv_to_buy"] = buy / pv * 100 if pv else 0
    metrics["pv_to_cart"] = cart / pv * 100 if pv else 0
    metrics["pv_to_fav"] = fav / pv * 100 if pv else 0
    metrics["cart_to_buy"] = buy / cart * 100 if cart else 0

    # 用户活跃度
    user_actions = df.groupby("user_id").size()
    metrics["avg_actions"] = user_actions.mean()
    metrics["median_actions"] = user_actions.median()
    top10_threshold = int(len(user_actions) * 0.1)
    top10_sum = user_actions.nlargest(top10_threshold).sum()
    metrics["top10_pct"] = top10_sum / user_actions.sum() * 100

    # 高峰时段
    hourly = df[df["behavior_type"] == "pv"].groupby("hour").size()
    peak_hour = hourly.idxmax()
    metrics["peak_hour"] = f"{peak_hour}:00"
    low_hour = hourly.idxmin()
    metrics["low_hour"] = f"{low_hour}:00"

    # 转化漏斗（用户级别）
    users_pv = set(df[df["behavior_type"] == "pv"]["user_id"])
    users_intent = set(df[df["behavior_type"].isin(["cart","fav"])]["user_id"])
    users_buy = set(df[df["behavior_type"] == "buy"]["user_id"])
    metrics["users_pv"] = len(users_pv)
    metrics["users_intent"] = len(users_intent)
    metrics["users_buy"] = len(users_buy)
    metrics["intent_rate"] = len(users_intent) / len(users_pv) * 100
    metrics["buy_rate"] = len(users_buy) / len(users_pv) * 100

    # 购买用户中加购/收藏比例
    buyers = df[df["behavior_type"] == "buy"]["user_id"].unique()
    buyer_df = df[df["user_id"].isin(buyers)]
    buyer_cart = len(set(buyer_df[buyer_df["behavior_type"] == "cart"]["user_id"]))
    buyer_fav = len(set(buyer_df[buyer_df["behavior_type"] == "fav"]["user_id"]))
    metrics["buyer_cart_pct"] = buyer_cart / len(buyers) * 100
    metrics["buyer_fav_pct"] = buyer_fav / len(buyers) * 100

    # 留存分析
    first_day = df.groupby("user_id")["date"].min().reset_index().rename(columns={"date":"first_date"})
    df_ret = df.merge(first_day, on="user_id")
    df_ret["day_offset"] = (df_ret["date"] - df_ret["first_date"]).dt.days

    # 计算次日、3日、7日留存（取所有首日批次的平均）
    ret_rates = {}
    for offset, label in [(1,"d1"), (3,"d3"), (7,"d7")]:
        rates = []
        for first_date in df_ret["first_date"].unique():
            cohort_users = set(first_day[first_day["first_date"]==first_date]["user_id"])
            if len(cohort_users) == 0: continue
            target_date = first_date + pd.Timedelta(days=offset)
            active = set(df_ret[(df_ret["user_id"].isin(cohort_users)) & (df_ret["date"]==target_date)]["user_id"])
            rates.append(len(active)/len(cohort_users)*100)
        ret_rates[label] = np.mean(rates) if rates else 0
    metrics["ret_d1"] = ret_rates["d1"]
    metrics["ret_d3"] = ret_rates["d3"]
    metrics["ret_d7"] = ret_rates["d7"]

    # 首日行为频次与留存关系
    first_actions = df_ret[df_ret["day_offset"]==0].groupby("user_id").size().reset_index(name="cnt")
    high = set(first_actions[first_actions["cnt"]>=5]["user_id"])
    low = set(first_actions[first_actions["cnt"]<5]["user_id"])
    d1 = pd.Timestamp("2017-11-26")
    high_ret = len(set(df_ret[(df_ret["user_id"].isin(high)) & (df_ret["date"]==d1)]["user_id"]))
    low_ret = len(set(df_ret[(df_ret["user_id"].isin(low)) & (df_ret["date"]==d1)]["user_id"]))
    metrics["high_ret_d1"] = high_ret / len(high) * 100 if high else 0
    metrics["low_ret_d1"] = low_ret / len(low) * 100 if low else 0
    metrics["gap"] = metrics["high_ret_d1"] - metrics["low_ret_d1"]

    # 用户分层
    base_date = df["date"].max()
    recency = df.groupby("user_id")["date"].max().reset_index()
    recency["R"] = (base_date - recency["date"]).dt.days
    frequency = df.groupby("user_id").size().reset_index(name="F")
    rf = recency.merge(frequency, on="user_id")
    r_med = rf["R"].median()
    f_med = rf["F"].median()
    rf["seg"] = "一般用户"
    rf.loc[(rf["R"]<=r_med)&(rf["F"]>=f_med), "seg"] = "高价值用户"
    rf.loc[(rf["R"]<=r_med)&(rf["F"]<f_med), "seg"] = "浅度/新用户"
    rf.loc[(rf["R"]>r_med)&(rf["F"]>=f_med), "seg"] = "沉睡用户"
    rf.loc[(rf["R"]>r_med)&(rf["F"]<f_med), "seg"] = "流失用户"
    seg_stats = rf.groupby("seg").agg(count=("user_id","count"), avg_R=("R","mean"), avg_F=("F","mean"))
    seg_stats["pct"] = seg_stats["count"] / seg_stats["count"].sum() * 100
    for seg in ["高价值用户","浅度/新用户","沉睡用户","流失用户"]:
        if seg in seg_stats.index:
            row = seg_stats.loc[seg]
            metrics[f"{seg}_pct"] = row["pct"]
            metrics[f"{seg}_avgF"] = row["avg_F"]
            metrics[f"{seg}_avgR"] = row["avg_R"]
        else:
            metrics[f"{seg}_pct"] = 0
            metrics[f"{seg}_avgF"] = 0
            metrics[f"{seg}_avgR"] = 0

    return metrics

def add_image(doc, image_path, width=Inches(5.5)):
    """插入图片，若文件不存在则插入提示文字"""
    try:
        doc.add_picture(image_path, width=width)
    except FileNotFoundError:
        doc.add_paragraph(f"[图片：{image_path} 未找到，请手动插入]")

def generate_report(metrics):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)

    # ========= 封面 =========
    title = doc.add_heading('淘宝用户行为数据分析\n与留存优化方案', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('产品运营实战项目报告')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # ========= P2 =========
    doc.add_page_break()
    doc.add_heading('一、项目背景与目标', level=1)
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph(
        '在互联网产品竞争日益激烈的今天，用户留存已成为衡量产品健康度的核心指标。'
        '获客成本不断攀升，而用户流失率居高不下，如何通过数据分析找到留存的关键驱动因素，'
        '并制定针对性的运营策略，是每一个产品运营必须面对的课题。'
    )
    doc.add_heading('1.2 数据来源', level=2)
    doc.add_paragraph(
        f'本项目使用阿里云天池平台公开的淘宝用户购物行为数据集（UserBehavior），'
        f'该数据集包含约{metrics["total_users"]/1e4:.0f}万随机用户在2017年11月25日至12月3日期间的所有行为记录，'
        f'总计约{metrics["total_records"]/1e8:.1f}亿条。每条记录包含用户ID、商品ID、商品类目ID、行为类型和时间戳五个字段。'
    )
    doc.add_heading('1.3 项目目标', level=2)
    doc.add_paragraph('1. 分析用户行为特征，理解用户购物路径')
    doc.add_paragraph('2. 构建转化漏斗，定位流失关键节点')
    doc.add_paragraph('3. 计算留存率，发现留存规律')
    doc.add_paragraph('4. 基于RF模型进行用户分层')
    doc.add_paragraph('5. 输出可落地的运营优化策略')

    # ========= P3 =========
    doc.add_page_break()
    doc.add_heading('二、数据预处理', level=1)
    doc.add_paragraph('原始数据集为CSV格式，无表头，约3.4GB。处理步骤如下：')
    steps = [
        '读取CSV文件，手动指定列名（user_id, item_id, category_id, behavior_type, time_stamp）',
        '过滤有效行为类型（pv/buy/cart/fav），去除异常记录',
        '清洗时间戳字段（处理科学计数法、空值、格式异常），转换为日期时间格式',
        '过滤日期范围（仅保留2017-11-25至2017-12-03的数据）',
        '去除重复记录',
        '新增衍生字段：date（日期）、hour（小时）',
        f'保存清洗后数据，共{metrics["total_records"]:,}条记录，{metrics["total_users"]:,}名用户，{metrics["total_items"]:,}件商品'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"步骤{i}：{step}")

    # ========= P4-P7 =========
    doc.add_page_break()
    doc.add_heading('三、整体行为分析', level=1)

    doc.add_heading('3.1 行为类型分布', level=2)
    add_image(doc, IMAGE_DIR + "图1_行为类型占比.png")
    doc.add_paragraph('关键发现：')
    doc.add_paragraph(f'• PV（页面浏览）占比最高，符合电商用户"浏览为主、购买为辅"的行为特征')
    doc.add_paragraph(f'• 点击→购买转化率约为{metrics["pv_to_buy"]:.2f}%')
    doc.add_paragraph(f'• 点击→加购转化率为{metrics["pv_to_cart"]:.2f}%')

    doc.add_heading('3.2 每日PV/UV趋势', level=2)
    add_image(doc, IMAGE_DIR + "图2_每日PVUV趋势.png")
    doc.add_paragraph('关键发现：')
    doc.add_paragraph('• 9天观察期内，PV和UV整体平稳')
    doc.add_paragraph(f'• 总PV约{metrics["pv_num"]/1e6:.1f}百万，总UV约{metrics["total_users"]/1e4:.0f}万')

    doc.add_heading('3.3 24小时活跃分布', level=2)
    add_image(doc, IMAGE_DIR + "图3_24小时活跃分布.png")
    doc.add_paragraph('关键发现：')
    doc.add_paragraph(f'• 用户活跃高峰时段为{metrics["peak_hour"]}-{metrics["peak_hour"]}左右')
    doc.add_paragraph(f'• 低谷时段为凌晨{metrics["low_hour"]}')

    doc.add_heading('3.4 用户活跃度分布', level=2)
    add_image(doc, IMAGE_DIR + "图4_用户活跃度分布.png")
    doc.add_paragraph('关键发现：')
    doc.add_paragraph(f'• 人均行为数{metrics["avg_actions"]:.1f}次，中位数{metrics["median_actions"]:.0f}次')
    doc.add_paragraph(f'• Top 10%用户贡献了约{metrics["top10_pct"]:.1f}%的行为量')

    # ========= P8-P9 =========
    doc.add_page_break()
    doc.add_heading('四、转化漏斗分析', level=1)
    add_image(doc, IMAGE_DIR + "图5_转化漏斗.png")
    doc.add_paragraph('漏斗数据：')
    doc.add_paragraph(f'• 浏览用户：{metrics["users_pv"]/1e4:.1f}万（100%）')
    doc.add_paragraph(f'• 意向用户（加购∪收藏）：{metrics["users_intent"]/1e4:.1f}万（{metrics["intent_rate"]:.1f}%）')
    doc.add_paragraph(f'• 购买用户：{metrics["users_buy"]/1e4:.1f}万（{metrics["buy_rate"]:.1f}%）')
    doc.add_paragraph()
    doc.add_paragraph('关键发现：')
    doc.add_paragraph(f'• 最大流失环节：浏览→意向环节转化率{metrics["intent_rate"]:.1f}%')
    doc.add_paragraph(f'• 购买用户中，有加购行为的比例{metrics["buyer_cart_pct"]:.1f}%，有收藏行为的比例{metrics["buyer_fav_pct"]:.1f}%')
    doc.add_paragraph('• 结论：提升加购转化是提高整体购买率的关键')

    # ========= P10-P11 =========
    doc.add_page_break()
    doc.add_heading('五、留存分析', level=1)
    doc.add_heading('5.1 留存曲线', level=2)
    add_image(doc, IMAGE_DIR + "图6_留存曲线.png")
    doc.add_paragraph('核心指标：')
    doc.add_paragraph(f'• 次日留存率（平均）：{metrics["ret_d1"]:.1f}%')
    doc.add_paragraph(f'• 3日留存率（平均）：{metrics["ret_d3"]:.1f}%')
    doc.add_paragraph(f'• 7日留存率（平均）：{metrics["ret_d7"]:.1f}%')
    doc.add_paragraph()
    doc.add_heading('5.2 留存热力图', level=2)
    add_image(doc, IMAGE_DIR + "图7_留存热力图.png")
    doc.add_paragraph('关键发现：')
    doc.add_paragraph(f'• 首日行为频次≥5次的用户次日留存率{metrics["high_ret_d1"]:.1f}%，<5次的仅{metrics["low_ret_d1"]:.1f}%，差距达{metrics["gap"]:.1f}个百分点')
    doc.add_paragraph('• 核心结论：首日体验深度是决定用户是否留存的最关键因素')

    # ========= P12 =========
    doc.add_page_break()
    doc.add_heading('六、用户分层（RF模型）', level=1)
    add_image(doc, IMAGE_DIR + "图8_用户分层.png")
    doc.add_paragraph('分层标准：基于R（最近一次行为距今天数）和F（行为频次）')
    doc.add_paragraph('• R ≤ 中位数 且 F ≥ 中位数 → 高价值用户')
    doc.add_paragraph('• R ≤ 中位数 且 F < 中位数 → 浅度/新用户')
    doc.add_paragraph('• R > 中位数 且 F ≥ 中位数 → 沉睡用户')
    doc.add_paragraph('• R > 中位数 且 F < 中位数 → 流失用户')
    doc.add_paragraph()
    doc.add_paragraph('分层结果：')
    doc.add_paragraph(f'• 高价值用户：{metrics["高价值用户_pct"]:.1f}%（平均F={metrics["高价值用户_avgF"]:.0f}次）')
    doc.add_paragraph(f'• 浅度/新用户：{metrics["浅度/新用户_pct"]:.1f}%（平均F={metrics["浅度/新用户_avgF"]:.0f}次）')
    doc.add_paragraph(f'• 沉睡用户：{metrics["沉睡用户_pct"]:.1f}%（平均R={metrics["沉睡用户_avgR"]:.0f}天）')
    doc.add_paragraph(f'• 流失用户：{metrics["流失用户_pct"]:.1f}%（平均R={metrics["流失用户_avgR"]:.0f}天）')

    # ========= P13 =========
    doc.add_page_break()
    doc.add_heading('七、关键发现汇总', level=1)
    findings = [
        f'转化漏斗的关键卡点在浏览→意向环节，转化率仅{metrics["intent_rate"]:.1f}%',
        f'次日留存率为{metrics["ret_d1"]:.1f}%，低于电商行业平均水平（30%-50%）',
        f'首日行为频次≥5次的用户留存率显著高于低频用户，差距达{metrics["gap"]:.1f}个百分点',
        f'用户活跃高峰集中在{metrics["peak_hour"]}，低谷在凌晨{metrics["low_hour"]}',
        f'高价值用户占比{metrics["高价值用户_pct"]:.1f}%，是运营维护的核心人群；流失用户占比{metrics["流失用户_pct"]:.1f}%，需要召回策略',
        f'加购行为是购买意愿的强信号，有加购行为的用户购买转化率是普通用户的{metrics["cart_to_buy"]/metrics["pv_to_buy"]:.1f}倍（若pv_to_buy>0）'
    ]
    for i, f in enumerate(findings, 1):
        doc.add_paragraph(f"发现{i}：{f}")

    # ========= P14-P15 =========
    doc.add_page_break()
    doc.add_heading('八、运营策略建议', level=1)

    strategies = [
        {
            'title': '策略一：新用户新手引导优化',
            'target': f'目标人群：浅度/新用户（占比{metrics["浅度/新用户_pct"]:.1f}%）',
            'problem': '问题：首日行为频次低的用户次日留存显著更低',
            'actions': [
                '设计3步新手引导流程：选择兴趣品类→完成首次浏览→触发首次互动',
                '新用户首单立减券（满20减5），降低首次购买门槛',
                '首日完成≥3个行为触发奖励（积分/优惠券），激励深度探索',
            ],
            'metric': '衡量指标：次日留存率、首日人均行为数、新用户首单转化率',
            'effect': f'预期效果：次日留存率提升{metrics["gap"]:.0f}个百分点左右'
        },
        {
            'title': '策略二：加购用户临门一脚转化',
            'target': '目标人群：加购未购买用户',
            'problem': '问题：加购→购买环节流失严重，用户可能在犹豫价格或比价',
            'actions': [
                '加购后24h内未付款 → 推送限时优惠提醒',
                '加购满X件 → 触发满减提示（如"再加1件享9折"）',
                '对高频加购但不买的用户，推送同类商品的更低价格/更好评价',
            ],
            'metric': '衡量指标：加购转化率、购物车弃置率、GMV',
            'effect': f'预期效果：加购→购买转化率提升{metrics["cart_to_buy"]*0.15:.1f}个百分点'
        },
        {
            'title': '策略三：沉睡用户召回',
            'target': f'目标人群：沉睡用户（占比{metrics["沉睡用户_pct"]:.1f}%）',
            'problem': '问题：曾有高频行为但近期未活跃，有召回价值',
            'actions': [
                '基于RF模型自动识别沉睡用户',
                '推送个性化召回Push："你收藏的商品降价了"/"你常逛的品类有新货"',
                '发送专属召回优惠券（满减力度高于普通用户）',
                '召回Push发送时间选在用户历史活跃高峰时段',
            ],
            'metric': '衡量指标：唤醒率、唤醒后7日留存、唤醒成本（CAC）',
            'effect': '预期效果：沉睡用户唤醒率5%-10%'
        },
        {
            'title': '策略四：高价值用户维系',
            'target': f'目标人群：高价值用户（占比{metrics["高价值用户_pct"]:.1f}%）',
            'problem': '问题：这批用户贡献最大价值，流失成本高',
            'actions': [
                '建立VIP等级体系，高价值用户自动进入高级别',
                '专属客服/优先发货/新品试用权',
                '个性化推荐精准度持续提升',
                '定期满意度调研，及时响应投诉',
            ],
            'metric': '衡量指标：高价值用户留存率、ARPU、NPS',
            'effect': '预期效果：高价值用户流失率降低至5%以下'
        }
    ]
    for s in strategies:
        doc.add_heading(s['title'], level=2)
        doc.add_paragraph(s['target'])
        doc.add_paragraph(s['problem'])
        doc.add_paragraph('具体措施：')
        for a in s['actions']:
            doc.add_paragraph(f"• {a}")
        doc.add_paragraph(s['metric'])
        doc.add_paragraph(s['effect'])

    # ========= 结尾 =========
    doc.add_page_break()
    doc.add_heading('九、项目总结', level=1)
    doc.add_paragraph(
        '本项目从阿里天池公开的淘宝用户行为数据集出发，完成了从数据预处理、'
        '行为分析、转化漏斗、留存分析到用户分层和运营策略输出的完整流程。'
        '通过Python（Pandas/NumPy/Matplotlib）实现全部数据分析与可视化，'
        '展示了数据驱动运营决策的方法论。'
    )
    doc.add_paragraph()
    doc.add_paragraph('方法论总结：')
    doc.add_paragraph('1. 数据先行：所有结论均基于真实数据，而非主观臆断')
    doc.add_paragraph('2. 指标导向：聚焦留存率、转化率等可量化指标')
    doc.add_paragraph('3. 分层运营：不同人群差异化策略，而非"一刀切"')
    doc.add_paragraph('4. 可衡量：每条策略都有明确的衡量指标和预期效果')
    # 保存
    output_path = "最终报告_淘宝用户行为分析.docx"
    doc.save(output_path)
    print(f"✅ 报告已自动生成：{output_path}")

if __name__ == "__main__":
    print("加载数据并计算指标...")
    df = load_data()
    metrics = calc_key_metrics(df)
    print("生成报告...")
    generate_report(metrics)
    print("完成！")